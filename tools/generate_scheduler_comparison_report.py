from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "outputs" / "reports"
DOCX_PATH = OUT_DIR / "港口水域UAV-USV上层调度模型与启发式规则对比分析.docx"
SCENARIO_IMAGE = ROOT / "data" / "ports" / "shanghai_yangshan_v1" / "preview.png"


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = _load_metrics()
    doc = Document()
    _setup_document(doc)
    _build_report(doc, data)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


def _load_metrics() -> dict[str, Any]:
    base = ROOT / "outputs" / "port_inspection" / "shanghai_yangshan_v1"
    return {
        "ppo_initial": _read_json(base / "scheduler_rl" / "scheduler_summary.json"),
        "greedy_global": _read_json(base / "greedy_env_global_score_summary.json"),
        "greedy_legacy": _read_json(base / "greedy_env_legacy_order_summary.json"),
        "offline_baseline": _read_json(base / "baseline_summary.json"),
    }


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    _set_style_font(styles["Normal"], "SimSun", 11, "000000", after_pt=6, line=1.10)
    _set_style_font(styles["Title"], "SimSun", 20, "0B2545", after_pt=8, bold=True)
    _set_style_font(styles["Subtitle"], "SimSun", 10.5, "555555", after_pt=12)
    _set_style_font(styles["Heading 1"], "SimSun", 16, "2E74B5", before_pt=16, after_pt=8, bold=True)
    _set_style_font(styles["Heading 2"], "SimSun", 13, "2E74B5", before_pt=12, after_pt=6, bold=True)
    _set_style_font(styles["Heading 3"], "SimSun", 12, "1F4D78", before_pt=8, after_pt=4, bold=True)
    _set_style_font(styles["List Bullet"], "SimSun", 11, "000000", after_pt=8, line=1.167)
    _set_style_font(styles["List Number"], "SimSun", 11, "000000", after_pt=8, line=1.167)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("港口水域 UAV-USV 异构协同巡检调度模型阶段分析")
    _format_run(run, size=9, color="666666")


def _set_style_font(style, font_name: str, size_pt: float, color: str, before_pt: float = 0, after_pt: float = 0, line: float = 1.10, bold: bool = False) -> None:
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.paragraph_format.space_before = Pt(before_pt)
    style.paragraph_format.space_after = Pt(after_pt)
    style.paragraph_format.line_spacing = line


def _format_run(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.font.bold = bold


def _build_report(doc: Document, data: dict[str, Any]) -> None:
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("港口水域 UAV-USV 上层调度模型与启发式规则对比分析")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("周报材料 | 上海港洋山深水港栅格化巡检场景 | 2026年6月")

    _add_callout(
        doc,
        "核心结论",
        "当前上层调度模型已形成可复现实验闭环：统一环境下，初版 PPO 训练记录优于两个贪心规则；"
        "但新版地图与线任务定义刚完成修订，原 PPO checkpoint 需要重新训练后才能作为最终新版场景结果。"
        "对比分析显示，未完成全部任务主要来自有限续航、返航余量、面任务耗时和无补给机制共同造成的资源紧约束。"
    )

    _add_section_scenario(doc)
    _add_section_model(doc)
    _add_section_reward(doc)
    _add_section_comparison(doc, data)
    _add_section_findings(doc, data)
    _add_section_limits_next(doc)
    _add_appendix(doc)


def _add_section_scenario(doc: Document) -> None:
    doc.add_heading("一、当前实验场景概述", level=1)
    p = doc.add_paragraph()
    p.add_run("当前场景为上海港洋山深水港水域巡检栅格化研究示意图。").bold = True
    p.add_run(" 地图不是海图或 GIS 精确复刻，而是依据公开空间结构抽象出主集装箱码头、自动化码头、东海大桥进港方向、主深水航道、LNG/能源邻近水域、外锚地/待泊水域等关键要素。")

    _add_bullets(
        doc,
        [
            "任务规模：32 个巡检任务，包括 16 个点任务、8 个线任务、8 个面任务。",
            "平台规模：8 个异构 agent，包括 4 架 UAV 与 4 艘 USV。",
            "点任务：航标、异常停留、污染/漂浮物等单点确认。",
            "线任务：主航道、进出港航路、泊位前沿线、防波堤或安全边界邻近水域。",
            "面任务：港池、泊位前沿高风险水域、自动化码头外侧水域、外锚地等区域覆盖。",
        ],
    )

    if SCENARIO_IMAGE.exists():
        doc.add_paragraph("图 1  当前上海港洋山深水港栅格化巡检场景示意。", style=None)
        pic_p = doc.add_paragraph()
        pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_p.add_run().add_picture(str(SCENARIO_IMAGE), width=Inches(6.3))


def _add_section_model(doc: Document) -> None:
    doc.add_heading("二、当前上层调度模型", level=1)
    doc.add_paragraph(
        "当前实现的是上层任务调度模型，采用集中式单智能体 PPO baseline。"
        "模型的决策对象不是低层连续运动控制，而是在每个调度步选择“某个平台执行某个任务”，用于验证任务建模、异构约束、动作掩码和奖励函数是否可学习。"
    )

    rows = [
        ("状态空间", "平台状态、任务状态、动作可行性 mask 和全局统计量。当前观测维度为 718。"),
        ("动作空间", "离散动作 (platform, task_choice)，其中 task_choice=0 表示等待。8个平台、32个任务时动作维度为 8×33=264。"),
        ("异构约束", "UAV 速度快、适合点/线快速筛查；USV 续航长、贴近水面，适合面任务和水面确认。动作 mask 会过滤不兼容或能量不足的任务。"),
        ("环境步含义", "一个 step 表示一次上层调度决策，不是低层栅格移动步。max_steps=64 表示每个 episode 最多进行 64 次任务选择/等待决策。"),
        ("路径与能耗代理", "UAV 使用曼哈顿距离近似，USV 使用水域 A* 路径；任务执行后扣除能量，并保留返航安全余量。"),
    ]
    _add_table(doc, ["组成", "当前实现"], rows, [1800, 7560])

    _add_callout(
        doc,
        "为何先用 PPO 而非直接 MAPPO",
        "当前 PPO 是上层调度的可运行 baseline，用于先验证场景、任务、奖励和动作约束。"
        "从论文主模型角度看，MAPPO 仍然更适合作为后续正式模型，因为补给、并行动作和多平台资源分配具有明显多智能体协同特征。"
    )


def _add_section_reward(doc: Document) -> None:
    doc.add_heading("三、奖励函数设计", level=1)
    doc.add_paragraph(
        "奖励函数采用“完成收益 + 高风险及时完成奖励 - 路径/能耗/风险暴露/逾期/负载不均衡惩罚”的结构。"
        "其中风险暴露用于刻画高风险任务长期未巡检带来的累计安全压力。"
    )
    doc.add_paragraph(
        "单任务风险暴露定义为：E_i(t)=r_i·τ_i(t)，其中 r_i 为任务风险等级，τ_i(t) 为任务未巡检时间。"
        "该指标对应持续监测与多机器人巡逻文献中的 weighted latency / idleness 思想。"
    )
    rows = [
        ("完成奖励", "complete_reward × risk × priority", "鼓励完成任务，风险越高收益越大。"),
        ("高风险及时奖励", "risk≥3 且未超过 max_interval 时给 bonus", "鼓励高风险任务尽早完成。"),
        ("路径成本", "- path_cost × path_length/100", "抑制不必要绕行。"),
        ("能耗成本", "- energy_cost × energy", "避免过度消耗平台续航。"),
        ("风险暴露", "- exposure_cost × 平均风险暴露", "惩罚高风险任务长期未巡检。"),
        ("逾期惩罚", "- late_penalty × 平均逾期时间", "惩罚超过最大允许巡检间隔的任务。"),
        ("负载不均衡", "- imbalance_penalty × 平台负载差", "避免任务集中到少数平台。"),
        ("非法动作", "- invalid_penalty", "惩罚不可行调度。"),
    ]
    _add_table(doc, ["奖励项", "形式", "作用"], rows, [1700, 3100, 4560])


def _add_section_comparison(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("四、与启发式/贪心规则的对比", level=1)
    doc.add_paragraph(
        "本阶段同时保留三类对照：离线传统贪心、同环境 legacy_order 贪心、同环境 global_score 贪心。"
        "其中同环境贪心与 PPO 使用相同能量扣除、返航余量、动作 mask 和 episode 长度，具有更强可比性。"
    )

    ppo = data["ppo_initial"]
    legacy = data["greedy_legacy"]
    global_score = data["greedy_global"]
    offline = data["offline_baseline"]

    rows = [
        (
            "离线传统贪心",
            "32/32",
            "不统计",
            "不统计",
            str(offline["total_path_length"]),
            "不扣能量",
            "非同口径：一次性分配，未完整承受能量与动作 mask 约束。",
        ),
        (
            "同环境 legacy_order 贪心",
            f'{legacy["completed_tasks"]}/{legacy["task_count"]}',
            str(legacy["late_tasks"]),
            f'{legacy["risk_exposure_sum"]:.0f}',
            str(legacy["total_path_length"]),
            f'{legacy["total_energy"]:.3f}',
            "按风险排序逐个选平台，进入 RL 环境后无法完成全部任务。",
        ),
        (
            "同环境 global_score 贪心",
            f'{global_score["completed_tasks"]}/{global_score["task_count"]}',
            str(global_score["late_tasks"]),
            f'{global_score["risk_exposure_sum"]:.0f}',
            str(global_score["total_path_length"]),
            f'{global_score["total_energy"]:.3f}',
            "每步选择当前最高分合法平台-任务对；统一环境下略优于 legacy_order。",
        ),
        (
            "初版 PPO 训练记录",
            f'{ppo["completed_tasks"]}/32',
            str(ppo["late_tasks"]),
            f'{ppo["risk_exposure_sum"]:.0f}',
            str(ppo["total_path_length"]),
            f'{ppo["total_energy"]:.3f}',
            "初版训练共 1,000,000 steps；由于地图与线任务近期已更新，需重新训练后作为新版最终结果。",
        ),
    ]
    _add_table(doc, ["方法", "完成任务", "逾期数", "风险暴露", "路径长度", "能耗", "说明"], rows, [1500, 900, 700, 900, 900, 800, 3660], font_size=8.5)

    _add_callout(
        doc,
        "对比口径说明",
        "离线传统贪心能够给出 32/32 的任务分配，但它没有像 RL 环境一样逐步扣除能量，也没有因为返航余量不足而过滤后续动作。"
        "因此它适合作为传统分配参考，不适合直接与 PPO 的 episode 结果比较。真正可比的是同环境贪心。"
    )


def _add_section_findings(doc: Document, data: dict[str, Any]) -> None:
    doc.add_heading("五、阶段性分析结论", level=1)
    ppo = data["ppo_initial"]
    global_score = data["greedy_global"]
    legacy = data["greedy_legacy"]
    points = [
        f"统一环境下，global_score 贪心完成 {global_score['completed_tasks']}/32，legacy_order 贪心完成 {legacy['completed_tasks']}/32，说明资源约束下规则方法也难以全覆盖。",
        f"初版 PPO 训练记录完成 {ppo['completed_tasks']}/32，风险暴露 {ppo['risk_exposure_sum']:.0f}，逾期任务 {ppo['late_tasks']} 个，优于两个同环境贪心结果。",
        "未完成任务集中在面任务，说明瓶颈主要出现在 USV 有效作业资源、面任务服务时间和能源约束，而不是简单的任务数量不足。",
        "当前环境没有补给、换电或回港充电动作，因此模型实际上求解的是一次出航窗口内的调度，而非完整持续巡检。",
        "离线贪心与 RL 结果差异暴露出一个重要建模事实：若忽略能耗和返航约束，任务覆盖率会被高估。",
    ]
    _add_bullets(doc, points)

    rows = [
        ("优势", "当前模型已经能体现异构平台、风险优先、能耗约束和动作可行性过滤，且 PPO 在统一环境中表现优于规则贪心。"),
        ("不足", "仍为单调度中心 PPO，不能充分表达多平台并行动作；无补给机制；新版场景调整后需要重新训练。"),
        ("解释重点", "完成率不足不应简单归因于算法无效，更可能是资源紧约束与无补给机制共同造成。"),
    ]
    _add_table(doc, ["判断维度", "分析"], rows, [1600, 7760])


def _add_section_limits_next(doc: Document) -> None:
    doc.add_heading("六、后续改进建议", level=1)
    doc.add_paragraph("建议按“先稳定对照，再扩展模型”的顺序推进。")
    _add_numbered(
        doc,
        [
            "基于最新版上海洋山场景重新训练 PPO，获得新版场景的正式训练曲线和最终 checkpoint。",
            "增加统一评估脚本，固定输出 PPO、随机策略、legacy_order、global_score、离线贪心的同口径对比表。",
            "引入回港补给/换电/充电动作，将一次出航调度扩展为多航次持续巡检调度。",
            "将上层模型从集中式 PPO 扩展到 MAPPO，使每个 UAV/USV 可以进行并行决策与协同资源分配。",
            "对奖励函数增加终局未完成任务惩罚和面任务优先约束，缓解模型过早消耗 USV 资源的问题。",
        ],
    )

    _add_callout(
        doc,
        "周报建议表述",
        "本周已完成上海港洋山深水港栅格化巡检场景、点线面任务定义、UAV-USV 异构参数、上层 PPO 调度 baseline、"
        "以及同环境贪心对照实验。初步结果表明，在考虑能耗与返航约束后，传统贪心无法完成全部任务，PPO 具备更优的风险暴露控制能力；"
        "下一步将基于新版场景重新训练，并引入 MAPPO 与补给机制。"
    )


def _add_appendix(doc: Document) -> None:
    doc.add_heading("附录：当前实验文件与复现实验命令", level=1)
    rows = [
        ("场景配置", "configs/port_shanghai_yangshan_v1.toml"),
        ("平台参数", "configs/platform_profiles_cn_common.toml"),
        ("地图数据", "data/ports/shanghai_yangshan_v1/shanghai_yangshan_v1_grid.json"),
        ("任务数据", "data/ports/shanghai_yangshan_v1/shanghai_yangshan_v1_tasks.json"),
        ("PPO训练脚本", "tools/train_port_scheduler_rl.py"),
        ("同环境贪心评估", "tools/evaluate_port_scheduler_greedy.py"),
    ]
    _add_table(doc, ["项目", "路径"], rows, [1800, 7560])
    _add_bullets(
        doc,
        [
            "渲染场景图：python tools/render_port_scenario.py --config configs/port_shanghai_yangshan_v1.toml",
            "检查调度环境：python tools/check_port_inspection_env.py --config configs/port_shanghai_yangshan_v1.toml --steps 10",
            "训练 PPO：python tools/train_port_scheduler_rl.py --config configs/port_shanghai_yangshan_v1.toml --steps 1000000",
            "评估同环境贪心：python tools/evaluate_port_scheduler_greedy.py --config configs/port_shanghai_yangshan_v1.toml --strategy global_score",
        ],
    )


def _add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, [9360])
    cell = table.cell(0, 0)
    _set_cell_shading(cell, "F4F6F9")
    _set_cell_margins(cell, 120, 120, 120, 120)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    _format_run(r, size=10.5, color="1F3A5F", bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    _format_run(r2, size=10.5)
    doc.add_paragraph()


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(item)


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.add_run(item)


def _add_table(doc: Document, headers: list[str], rows: list[tuple[Any, ...]], widths: list[int], font_size: float = 9.5) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_shading(cell, "F2F4F7")
        _set_cell_margins(cell, 80, 80, 120, 120)
        _set_cell_text(cell, str(header), bold=True, size=font_size, color="0B2545", center=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            _set_cell_margins(cell, 80, 80, 120, 120)
            _set_cell_text(cell, str(value), size=font_size, center=False)
    doc.add_paragraph()


def _set_cell_text(cell, text: str, bold: bool = False, size: float = 9.5, color: str = "000000", center: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _format_run(run, size=size, color=color, bold=bold)


def _set_table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    tbl_pr.append(tbl_ind)
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[index]))


def _set_cell_margins(cell, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


if __name__ == "__main__":
    main()
