from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
OUT_PATH = OUT_DIR / "CIR_GAT_MAPPO_paper_outline_cn.docx"


def set_east_asia(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    set_east_asia(run, "宋体")
    run.font.size = Pt(9.5)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)
    for table_row in table.rows:
        for cell in table_row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.15
                paragraph.paragraph_format.space_after = Pt(1)
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_number(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.add_run(text)


def add_note(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF2CC")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title)
    set_east_asia(run, "黑体")
    run.bold = True
    run.font.size = Pt(10.5)
    paragraph.add_run("\n" + body)
    doc.add_paragraph()


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(3)

    for style_name, font_name, size in [
        ("Heading 1", "黑体", 15),
        ("Heading 2", "黑体", 12.5),
        ("Heading 3", "黑体", 11),
    ]:
        style = doc.styles[style_name]
        style.font.name = font_name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build() -> Path:
    doc = Document()
    configure_styles(doc)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("CIR-GAT-MAPPO中文小论文构架")
    set_east_asia(run, "黑体")
    run.font.size = Pt(22)
    run.bold = True

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("面向灾后未知危险区域覆盖勘察的多机器人协同覆盖方法")
    set_east_asia(run, "黑体")
    run.font.size = Pt(14)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("论文框架草案 · 按中文工科小论文写作风格整理")
    run.font.size = Pt(10.5)
    run.italic = True

    add_note(
        doc,
        "应用场景边界说明",
        "本文建议采用“灾后未知危险区域覆盖勘察/态势获取”作为应用场景，而不是完整灾后搜救。当前模型不包含搜救目标、幸存者概率矩阵、目标发现概率或任务优先级推断，因此不能声称解决完整搜救决策问题。更准确的表述是：机器人集群在灾后或危险未知区域中进行快速全覆盖巡查、环境勘察和搜索前态势获取，为后续人工救援、目标识别或路径规划提供覆盖信息。",
    )

    doc.add_heading("0 写作定位与推荐应用场景", level=1)
    doc.add_paragraph(
        "推荐应用场景：灾后未知危险区域的多机器人快速覆盖勘察。该场景可包括震后建筑内部、厂区事故现场、坍塌通道、地下空间、仓储火灾后区域等。任务目标不是直接定位救援目标，而是在局部可观测、通信受限和地图未知条件下，尽快覆盖可通行区域，减少重复搜索，为后续救援判断提供环境态势。"
    )
    add_bullet(doc, "为什么不写“完整灾后搜救”：当前模型没有目标概率地图、目标检测传感器模型、幸存者状态估计或救援任务分配。")
    add_bullet(doc, "为什么仍可写灾后场景：未知环境、障碍物、局部观测、通信半径、快速覆盖和减少重复搜索，均与灾后前期勘察任务高度契合。")
    add_bullet(doc, "若面向兵工类期刊，可将场景扩展为“危险未知区域无人集群侦察覆盖”；若面向农机/机器人类期刊，可表述为“未知作业区域多机器人覆盖巡检”。")

    add_table(
        doc,
        ["候选场景", "契合度", "推荐程度", "说明"],
        [
            ["灾后未知危险区域覆盖勘察", "高", "首选", "不需要目标概率矩阵，强调覆盖、建图和态势获取。"],
            ["未知环境无人集群侦察", "高", "可选", "适合偏兵工或无人系统方向，但注意避免过度军事化表述。"],
            ["远洋捕鱼未知区域覆盖", "中", "不推荐作为主场景", "捕鱼更依赖资源概率、目标迁移和收益模型，当前模型不包含。"],
            ["仓储/厂区巡检覆盖", "中高", "备选", "工程稳妥，但灾后场景的问题紧迫性更强。"],
        ],
    )

    doc.add_heading("1 题目与摘要设计", level=1)
    doc.add_heading("1.1 题目建议", level=2)
    add_number(doc, "面向灾后未知危险区域覆盖勘察的覆盖意图关系感知GAT-MAPPO方法")
    add_number(doc, "覆盖意图冲突感知的未知区域多机器人协同覆盖路径规划方法")
    add_number(doc, "基于边关系图注意力MAPPO的多机器人未知环境覆盖勘察方法")
    doc.add_paragraph("推荐使用第1个题目。该题目同时体现应用背景、核心任务和方法创新，且不会暗示模型完成目标搜救或概率推断。")

    doc.add_heading("1.2 摘要草稿", level=2)
    doc.add_paragraph(
        "针对灾后危险区域前期勘察中未知环境覆盖效率低、机器人间重复搜索严重的问题，提出一种覆盖意图关系感知的图注意力多智能体近端策略优化方法。该方法面向局部观测和有限通信条件下的多机器人协同覆盖任务，在集中训练、分散执行框架中引入显式局部地图记忆和覆盖意图消息。为增强通信邻居之间未来覆盖任务冲突的表达能力，设计覆盖意图关系模块，利用邻居间未来覆盖意图的Soft-IoU计算重叠强度，并将其作为零初始化、带界可学习边偏置注入图注意力分数。该模块不改变动作空间、奖励函数和通信拓扑，能够在原GAT-MAPPO基础上实现单变量扩展。实验在未知栅格环境中开展，并与无通信MAPPO、GAT-MAPPO及继续训练对照模型进行比较。结果表明，所提方法在保持覆盖完成率的同时提高了Coverage-AUC和早期覆盖率，并降低重复覆盖比例，验证了覆盖意图关系建模对未知区域覆盖勘察任务的有效性。"
    )

    doc.add_heading("1.3 关键词", level=2)
    doc.add_paragraph("多机器人系统；未知环境覆盖；图注意力网络；多智能体强化学习；覆盖意图关系；灾后勘察")

    doc.add_heading("2 引言构架", level=1)
    add_number(doc, "背景引入：灾后建筑、危险厂区和地下空间往往存在地图未知、障碍物不确定和通信受限等问题，需要多机器人快速完成区域覆盖勘察。")
    add_number(doc, "问题矛盾：多机器人能够提高覆盖速度，但由于局部观测和信息不完全，机器人间容易发生重复搜索、局部拥挤和后期补漏效率低等问题。")
    add_number(doc, "现有方法不足：传统覆盖路径规划依赖完整地图或规则分区；MAPPO具备协同学习能力，但普通MAPPO缺少显式通信结构；GAT-MAPPO能聚合邻居消息，但通常将未来覆盖意图作为节点属性，未显式建模“谁与我的未来覆盖任务重叠”。")
    add_number(doc, "本文思路：在现有GAT-MAPPO基础上，将覆盖意图重叠提升为边关系，以关系偏置形式注入GAT注意力分数，使策略能够选择性关注潜在冲突邻居。")

    add_table(
        doc,
        ["贡献点", "论文表述建议", "注意边界"],
        [
            ["MCPP任务适配", "构建面向未知区域覆盖勘察的GAT-MAPPO框架。", "不要声称解决完整搜救。"],
            ["覆盖意图关系", "提出基于Soft-IoU的邻居未来覆盖意图重叠关系。", "关系来自已通信意图，不读取通信范围外真值。"],
            ["零初始化边偏置", "设计带界可学习β，使模型初始退化为原GAT。", "不要夸大为理论最优保证。"],
            ["实验验证", "通过AUC、T90/T95、RepeatRatio等指标验证覆盖效率和重复率改善。", "报告同随机地图与训练seed差异。"],
        ],
    )

    doc.add_heading("3 问题建模", level=1)
    doc.add_heading("3.1 未知区域覆盖勘察任务", level=2)
    doc.add_paragraph(
        "将待勘察区域离散为二维栅格地图。机器人初始不知道完整地图，只能通过局部观测逐步获得周围空闲栅格、障碍物和已覆盖信息。覆盖定义为机器人实际占据过的可通行栅格，而不是传感器视野扫过的区域。任务目标是在最大步数内尽可能覆盖全部可通行栅格，并减少重复访问。"
    )

    doc.add_heading("3.2 观测、通信与动作", level=2)
    add_bullet(doc, "观测：每个机器人获得局部窗口观测、局部地图记忆、覆盖进度和任务阶段元数据。")
    add_bullet(doc, "通信：机器人仅与通信半径内邻居交换覆盖消息，通信图随机器人位置动态变化。")
    add_bullet(doc, "动作：采用上、下、左、右、停止等离散动作，非法动作通过动作掩码屏蔽。")
    add_bullet(doc, "训练范式：采用集中训练、分散执行；actor不使用执行阶段不可获得的全局覆盖真值。")

    add_table(
        doc,
        ["符号", "含义"],
        [
            ["N", "机器人数量"],
            ["A_t", "t时刻通信邻接矩阵"],
            ["o_i", "机器人i的局部观测与地图记忆"],
            ["m_i", "机器人i的覆盖意图消息"],
            ["I_i", "机器人i的未来覆盖意图表示"],
            ["rho_ij", "机器人i与j的意图重叠强度"],
            ["e_ij", "原始GAT注意力logit"],
            ["beta", "可学习覆盖意图关系偏置系数"],
        ],
    )

    doc.add_heading("4 CIR-GAT-MAPPO方法构架", level=1)
    doc.add_heading("4.1 总体框架", level=2)
    doc.add_paragraph(
        "方法采用MAPPO作为基础训练框架。Actor根据本体局部观测、邻居覆盖消息和动态图注意力输出动作分布；Critic使用集中训练阶段的全局栅格状态估计价值函数。CIR模块只作用于actor侧真实邻居聚合GAT，不修改critic、奖励函数和PPO损失。"
    )

    doc.add_heading("4.2 显式局部地图记忆与覆盖消息", level=2)
    doc.add_paragraph(
        "机器人维护已知空闲区域、已知障碍物、团队已覆盖区域、未知区域和frontier集合。覆盖消息包含已知覆盖率、自身覆盖率、未知区域比例、frontier比例、近期新覆盖率、近期重复率、停滞程度、目标方向、相对目标位置、意图区域和意图有效标志。"
    )

    doc.add_heading("4.3 边特征图注意力通信", level=2)
    doc.add_paragraph(
        "通信图由通信半径动态决定。GAT在通信硬掩码内计算邻居注意力，并结合相对距离、相对行列偏移和连通标志等边特征。相较无通信MAPPO，该结构能够根据邻近机器人空间关系和覆盖消息进行协同决策。"
    )

    doc.add_heading("4.4 覆盖意图关系计算", level=2)
    doc.add_paragraph("将每个机器人未来覆盖意图表示为统一意图区域向量。对于机器人i和j，采用Soft-IoU计算二者意图重叠：")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("inter_ij = Σ_x I_i(x) I_j(x),   union_ij = Σ_x I_i(x)+Σ_x I_j(x)-inter_ij,   ρ_ij = inter_ij/(union_ij+ε)")
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    doc.add_paragraph("其中，ρij∈[0,1]，数值越大表示两个机器人未来覆盖意图越重叠。自环关系置为0，非通信边关系置为0并继续由GAT硬掩码屏蔽。")

    doc.add_heading("4.5 CIR边注意力偏置", level=2)
    doc.add_paragraph("在原始GAT注意力logit上加入覆盖意图关系偏置：")
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("e'_ij = e_ij + βρ_ij,       β = βmax · tanh(b)")
    run.font.name = "Cambria Math"
    run.font.size = Pt(10)
    doc.add_paragraph(
        "其中b初始化为0，因此初始β=0，模型严格退化为原始GAT-MAPPO，便于从已有GAT-ON模型进行warm-start，也降低训练早期扰动。β可学习为正或负，由策略自行决定是增强还是抑制高冲突邻居的注意力。"
    )
    add_note(doc, "方法边界", "CIR不增加通信字段，不改变通信范围，不修改奖励函数，不加入agent ID，不加入额外辅助损失。其唯一核心改动是在actor侧GAT softmax之前加入βρij。")

    doc.add_heading("5 实验设计", level=1)
    doc.add_heading("5.1 实验环境", level=2)
    add_bullet(doc, "地图：二维未知栅格环境，主测试场景为20×20。")
    add_bullet(doc, "机器人数量：4个。")
    add_bullet(doc, "障碍物：随机障碍物比例约5%，测试时使用固定随机地图种子。")
    add_bullet(doc, "观测与通信：局部观测半径2，通信半径4，使用显式局部地图记忆与覆盖消息。")
    add_bullet(doc, "评价范围：只评价覆盖阶段，不引入返回阶段，避免返回策略噪声影响覆盖结论。")

    doc.add_heading("5.2 对比方法", level=2)
    add_table(
        doc,
        ["模型", "说明", "作用"],
        [
            ["GAT-OFF / MAPPO", "不使用图注意力通信或关闭通信增强。", "验证通信模块必要性。"],
            ["GAT-ON / GAT-MAPPO", "使用覆盖消息、边特征和动态图注意力。", "主要基线模型。"],
            ["GAT-ON-continued", "从GAT-ON继续训练相同步数但不开CIR。", "排除“继续训练导致提升”的干扰。"],
            ["CIR-GAT-MAPPO", "在GAT-ON基础上加入覆盖意图关系边偏置。", "本文方法。"],
        ],
    )

    doc.add_heading("5.3 评价指标", level=2)
    add_table(
        doc,
        ["指标", "含义", "方向"],
        [
            ["Coverage-AUC", "覆盖率-时间曲线面积，衡量整体覆盖效率。", "越大越好"],
            ["T90/T95", "达到90%/95%覆盖率所需步数。", "越小越好"],
            ["Coverage@100/200/300", "固定步数下的覆盖率。", "越大越好"],
            ["RepeatRatio", "团队重复访问比例。", "越小越好"],
            ["Success Rate", "最大步数内完成覆盖的比例。", "越大越好"],
            ["Steps/PathLength", "完成覆盖所需步数和总路径长度。", "越小越好"],
            ["IntentConflictRate", "通信边上的平均意图重叠强度。", "机制诊断"],
            ["β与attention entropy", "CIR使用关系偏置的程度和注意力集中程度。", "机制解释"],
        ],
    )

    doc.add_heading("6 结果分析写法建议", level=1)
    doc.add_heading("6.1 主结果表", level=2)
    doc.add_paragraph(
        "主结果表建议报告同随机地图20个seed的均值，并突出CIR-policy相对GAT-ON与GAT-ON-continued的变化。根据当前实验，CIR-policy在Coverage-AUC、T90、Coverage@100、RepeatRatio和平均Steps上优于GAT-ON，同时显著优于GAT-ON-continued控制组。"
    )
    add_table(
        doc,
        ["模型", "AUC", "T90", "T95", "Cov@100", "RepeatRatio", "完成率", "Steps"],
        [
            ["GAT-ON", "0.8646", "131.5", "135.1", "79.17%", "46.85%", "95%", "191.7"],
            ["GAT-ON-continued-last", "0.8629", "136.6", "152.3", "77.09%", "45.00%", "95%", "186.8"],
            ["CIR-GAT-MAPPO", "0.8677", "130.5", "147.4", "80.50%", "42.47%", "95%", "179.7"],
        ],
    )

    doc.add_heading("6.2 建议表述", level=2)
    doc.add_paragraph(
        "可以写：与GAT-MAPPO相比，CIR-GAT-MAPPO在相同随机测试地图上提高Coverage-AUC，并降低重复覆盖率和平均完成步数；与继续训练但不开CIR的控制组相比，CIR在多数seed上保持更好的早期覆盖效率和更低重复率，说明性能提升并非简单由额外训练步数造成。"
    )

    doc.add_heading("6.3 需要谨慎表述的现象", level=2)
    add_bullet(doc, "当前最佳CIR checkpoint中β较小，说明关系偏置使用较保守，因此不要写成“强制冲突规避”。")
    add_bullet(doc, "CIR在训练seed上不一定最优，说明其优势更体现在随机地图泛化，而不是训练地图拟合。")
    add_bullet(doc, "T95均值并非所有对比中最优，可解释为中后段补漏仍存在波动，后续可改进意图时域或更精细路径意图表达。")

    doc.add_heading("7 图表与论文版式建议", level=1)
    add_table(
        doc,
        ["图/表编号", "建议内容", "目的"],
        [
            ["图1", "灾后未知危险区域覆盖勘察任务示意图。", "说明场景但避免目标搜救概率建模。"],
            ["图2", "CIR-GAT-MAPPO总体框架。", "展示局部观测、覆盖消息、GAT、CIR和PPO。"],
            ["图3", "覆盖意图关系ρij计算示意。", "解释Soft-IoU边关系。"],
            ["图4", "覆盖率曲线。", "展示Coverage-AUC和早期覆盖效率。"],
            ["表1", "实验参数设置。", "增强可复现性。"],
            ["表2", "主性能对比。", "报告AUC、T90/T95、RepeatRatio等。"],
            ["表3", "消融实验。", "比较GAT-ON、continued和CIR。"],
            ["表4", "机制指标。", "报告β、IntentConflictRate和attention entropy。"],
        ],
    )

    doc.add_heading("8 结论草稿", level=1)
    doc.add_paragraph(
        "本文面向灾后未知危险区域覆盖勘察任务，提出一种覆盖意图关系感知的CIR-GAT-MAPPO方法。该方法在多机器人GAT-MAPPO框架中显式构建通信邻居之间的未来覆盖意图重叠关系，并以零初始化、带界可学习边偏置注入actor侧图注意力分数。实验结果表明，所提方法在保持覆盖完成率的同时提升早期覆盖效率，降低重复覆盖率和平均完成步数。后续工作将进一步引入更精细的未来路径意图表示，并扩展至带目标检测或概率地图的完整搜救任务。"
    )

    doc.add_heading("9 写作红线与推荐措辞", level=1)
    add_table(
        doc,
        ["不要这样写", "建议这样写"],
        [
            ["解决灾后搜救目标定位问题", "服务于灾后未知区域快速覆盖勘察与态势获取"],
            ["根据幸存者概率矩阵规划路径", "根据局部地图记忆和覆盖意图进行协同覆盖"],
            ["完成完整搜救闭环", "为后续目标识别、人工救援和路径规划提供覆盖信息"],
            ["CIR强制避免冲突", "CIR显式提供意图重叠关系，由策略学习如何利用"],
            ["大幅全面优于所有指标", "在AUC、早期覆盖和重复率等指标上取得改善"],
        ],
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("生成日期：2026-06-08")
    run.font.size = Pt(9)
    run.italic = True

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
